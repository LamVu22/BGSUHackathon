#!/usr/bin/env python3
"""Extract clean text and link structure from crawled files.

Reads data/raw/metadata.tsv (produced by the crawler), parses HTML plus
non-HTML assets (PDF, DOCX, spreadsheets, etc.), and emits
data/processed/clean_nodes.json + clean_edges.json. These files are later
consumed by scripts/build_graph.py to compute graph metrics.

Usage (run from repo root with venv activated):

    python scripts/clean_content.py

Configuration lives in config/pipeline.json (or override with PIPELINE_CONFIG).
"""

from __future__ import annotations

import json
import logging
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Tuple
from urllib.parse import urljoin, urlparse

from bs4 import BeautifulSoup

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dep
    Document = None  # type: ignore

try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover
    fitz = None  # type: ignore


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_CONFIG_PATH = REPO_ROOT / "config" / "pipeline.json"


@dataclass
class CleaningSettings:
    raw_output: Path = REPO_ROOT / "data/raw"
    processed_output: Path = REPO_ROOT / "data/processed"
    allowed_domains: List[str] = field(default_factory=lambda: ["www.bgsu.edu", "bgsu.edu"])
    root_url: str = "https://www.bgsu.edu"
    snippet_chars: int = 600
    checkpoint_interval: int = 50

    @classmethod
    def from_dict(cls, data: Dict[str, object]) -> "CleaningSettings":
        default = cls()

        def get_value(key: str, fallback):
            return data.get(key, fallback) if isinstance(data, dict) else fallback

        raw_output = _resolve_path(get_value("raw_output", default.raw_output))
        processed_output = _resolve_path(get_value("processed_output", default.processed_output))

        return cls(
            raw_output=raw_output,
            processed_output=processed_output,
            allowed_domains=list(get_value("allowed_domains", list(default.allowed_domains))),
            root_url=get_value("root_url", default.root_url),
            snippet_chars=int(get_value("graph_snippet_chars", default.snippet_chars)),
            checkpoint_interval=int(get_value("cleaning_checkpoint_interval", default.checkpoint_interval)),
        )

    @property
    def metadata_path(self) -> Path:
        return self.raw_output / "metadata.tsv"

    @property
    def clean_nodes_path(self) -> Path:
        return self.processed_output / "clean_nodes.json"

    @property
    def clean_edges_path(self) -> Path:
        return self.processed_output / "clean_edges.json"


class ContentCleaner:
    def __init__(self, settings: CleaningSettings) -> None:
        self.settings = settings
        self.nodes: Dict[str, Dict] = {}
        self.edges: List[Dict] = []

    def run(self) -> None:
        records = self._read_metadata()
        if not records:
            logging.error("No metadata records found at %s", self.settings.metadata_path)
            return

        if self.settings.clean_nodes_path.exists():
            try:
                with self.settings.clean_nodes_path.open("r", encoding="utf-8") as f:
                    existing_nodes = {node["url"].rstrip("/"): node for node in json.load(f)}
                self.nodes.update(existing_nodes)
                logging.info("Loaded %s existing cleaned nodes", len(self.nodes))
            except Exception as exc:
                logging.warning("Failed to load existing cleaned nodes: %s", exc)

        if self.settings.clean_edges_path.exists():
            try:
                with self.settings.clean_edges_path.open("r", encoding="utf-8") as f:
                    self.edges = json.load(f)
                logging.info("Loaded %s existing cleaned edges", len(self.edges))
            except Exception as exc:
                logging.warning("Failed to load existing cleaned edges: %s", exc)

        total = len(records)
        logging.info("Cleaning %s records (serial)", total)
        for idx, record in enumerate(records, start=1):
            url, node, edges = self._process_record(record)
            self.nodes[url] = node
            self.edges.extend(edges)
            if idx % 50 == 0:
                logging.info("Processing [%s/%s]: %s", idx, total, url)

        self._write_outputs()

    def _write_outputs(self, partial: bool = False) -> None:
        output_dir = self.settings.processed_output
        output_dir.mkdir(parents=True, exist_ok=True)
        with self.settings.clean_nodes_path.open("w", encoding="utf-8") as f:
            json.dump(list(self.nodes.values()), f, indent=2)
        with self.settings.clean_edges_path.open("w", encoding="utf-8") as f:
            json.dump(self.edges, f, indent=2)
        if partial:
            logging.info(
                "Checkpoint: %s nodes / %s edges written",
                len(self.nodes),
                len(self.edges),
            )
        else:
            logging.info(
                "Wrote %s nodes and %s edges to %s / %s",
                len(self.nodes),
                len(self.edges),
                self.settings.clean_nodes_path,
                self.settings.clean_edges_path,
            )

    def _process_record(self, record: Dict[str, str]):
        url = record["url"].rstrip("/")
        absolute_path = _resolve_path(record["path"])
        parsed = urlparse(url)
        node = {
            "url": url,
            "path": record["path"],
            "content_type": record["content_type"],
            "doc_type": self._infer_doc_type(record["path"], record["content_type"]),
            "title": None,
            "word_count": 0,
            "clean_text": "",
            "snippet": "",
            "domain": parsed.netloc,
            "is_root": url.rstrip("/") == self.settings.root_url.rstrip("/"),
        }
        edges: List[Dict[str, str]] = []

        if self._is_html(record["content_type"], record["path"]):
            title, word_count, links, clean_text = self._process_html(absolute_path, url)
            node["title"] = title
            node["word_count"] = word_count
            node["clean_text"] = clean_text
            node["snippet"] = clean_text[: self.settings.snippet_chars]
            for link_url, anchor_text in links:
                if not self._is_allowed_domain(link_url):
                    continue
                edges.append({"source": url, "target": link_url.rstrip("/"), "anchor_text": anchor_text})
        else:
            extracted_text = self._extract_text_from_file(absolute_path)
            node["title"] = Path(record["path"]).name
            node["clean_text"] = extracted_text
            node["snippet"] = extracted_text[: self.settings.snippet_chars]
            node["word_count"] = len(extracted_text.split())

        return url, node, edges

    def _read_metadata(self) -> List[Dict[str, str]]:
        metadata_path = self.settings.metadata_path
        if not metadata_path.exists():
            logging.error("Metadata file missing: %s", metadata_path)
            return []
        records: List[Dict[str, str]] = []
        with metadata_path.open("r", encoding="utf-8") as meta:
            header = meta.readline()
            for line in meta:
                parts = line.strip().split("\t")
                if len(parts) != 3:
                    continue
                url, path_str, content_type = parts
                records.append({"url": url, "path": path_str, "content_type": content_type})
        return records

    @staticmethod
    def _infer_doc_type(path: str, content_type: str) -> str:
        extension = Path(path).suffix.lower()
        if extension:
            return extension.lstrip(".")
        if content_type:
            return content_type.split("/")[-1]
        return "unknown"

    @staticmethod
    def _is_html(content_type: str, path: str) -> bool:
        if "text/html" in content_type:
            return True
        return Path(path).suffix.lower() in {".html", ".htm", ".php", ".asp", ".aspx", ".jsp"}

    def _is_allowed_domain(self, url: str) -> bool:
        parsed = urlparse(url)
        return parsed.netloc in self.settings.allowed_domains

    def _process_html(self, path: Path, base_url: str) -> Tuple[str, int, List[Tuple[str, str]], str]:
        try:
            html = path.read_text(encoding="utf-8", errors="ignore")
        except FileNotFoundError:
            logging.warning("HTML file missing during cleaning: %s", path)
            return "", 0, [], ""

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript", "svg"]):
            tag.decompose()

        title = (soup.title.string or "").strip() if soup.title and soup.title.string else ""
        text = soup.get_text(separator=" ", strip=True)
        clean_text = " ".join(text.split())
        word_count = len(clean_text.split())

        links: List[Tuple[str, str]] = []
        for anchor in soup.find_all("a", href=True):
            href = anchor["href"].strip()
            absolute = urljoin(base_url, href).split("#", 1)[0].rstrip("/")
            if not absolute:
                continue
            anchor_text = anchor.get_text(strip=True)[:200]
            links.append((absolute, anchor_text))
        return title, word_count, links, clean_text

    def _extract_text_from_file(self, path: Path) -> str:
        if not path.exists():
            logging.warning("File missing during cleaning: %s", path)
            return ""
        suffix = path.suffix.lower()
        if suffix in {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".svg", ".tif", ".tiff"}:
            return ""
        try:
            if suffix == ".pdf":
                if fitz is None:
                    logging.warning("PyMuPDF not installed; skipping PDF %s", path)
                    return ""
                doc = fitz.open(path)
                texts = []
                for page in doc:
                    txt = page.get_text("text").strip()
                    if txt:
                        texts.append(txt)
                doc.close()
                return " ".join(texts)
            if suffix in {".docx"}:
                if Document is None:
                    logging.warning("python-docx not installed; skipping DOCX %s", path)
                    return ""
                document = Document(path)
                parts: List[str] = []
                for para in document.paragraphs:
                    txt = para.text.strip()
                    if txt:
                        parts.append(txt)
                for table in document.tables:
                    for row in table.rows:
                        row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            parts.append(row_text)
                return " ".join(parts)
            logging.debug("Skipping unsupported asset %s", path)
            return ""
        except Exception as exc:  # pragma: no cover
            logging.warning("Failed extracting text from %s: %s", path, exc)
            return ""


def _resolve_path(path_value) -> Path:
    path = path_value if isinstance(path_value, Path) else Path(path_value)
    if not path.is_absolute():
        path = REPO_ROOT / path
    return path


def load_settings(config_path: Path | None = None) -> CleaningSettings:
    path = config_path or DEFAULT_CONFIG_PATH
    if path.exists():
        try:
            with path.open("r", encoding="utf-8") as f:
                data = json.load(f)
            logging.info("Loaded pipeline config from %s", path)
            return CleaningSettings.from_dict(data)
        except json.JSONDecodeError as exc:
            logging.error("Failed to parse config %s: %s", path, exc)
    else:
        logging.warning("Config file %s not found. Using defaults.", path)
    return CleaningSettings()


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
    config_path = Path(os.environ["PIPELINE_CONFIG"]) if "PIPELINE_CONFIG" in os.environ else None
    settings = load_settings(config_path)
    cleaner = ContentCleaner(settings)
    cleaner.run()


if __name__ == "__main__":
    main()
